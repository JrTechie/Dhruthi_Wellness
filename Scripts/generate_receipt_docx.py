import os
import sys
import datetime
import docx
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

# Define brand colors
COLOR_PRIMARY = RGBColor(0x3B, 0x53, 0x36)      # Forest Green #3B5336
COLOR_MUTED = RGBColor(0x5A, 0x6E, 0x58)        # Green Grey #5A6E58
COLOR_DARK = RGBColor(0x1C, 0x2B, 0x1A)         # Dark Charcoal #1C2B1A
COLOR_SUCCESS = RGBColor(0x2E, 0x7D, 0x32)      # Success Green #2E7D32

def set_cell_background(cell, fill_hex):
    """Sets background color of a Word table cell."""
    tcPr = cell._element.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    """Sets internal padding of a table cell in dxa units."""
    tcPr = cell._element.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{m}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def create_receipt_docx(
    output_filename="Dhruthi_Wellness_Receipt_Template.docx",
    receipt_no="DW-2026-001",
    payment_date=None,
    client_name="Ananya Sharma",
    client_phone="+91 98765 43210",
    client_email="client@example.com",
    plan_name="Dhruthi Elite — 3 Months (Diabetes / PCOS / Thyroid)",
    amount_paid=4899.00,
    payment_method="UPI (Google Pay)",
    transaction_id="TXN9876543210",
    notes="Payment proof verified. Thank you for choosing Dhruthi Wellness!"
):
    if payment_date is None:
        payment_date = datetime.date.today().strftime("%d %b, %Y")

    doc = Document()
    
    # Page setup: US Letter with 0.6 inch margins
    sections = doc.sections
    for section in sections:
        section.page_width = Inches(8.5)
        section.page_height = Inches(11.0)
        section.top_margin = Inches(0.6)
        section.bottom_margin = Inches(0.6)
        section.left_margin = Inches(0.6)
        section.right_margin = Inches(0.6)

    # Base Normal Style
    normal_style = doc.styles['Normal']
    normal_style.font.name = 'Arial'
    normal_style.font.size = Pt(10)
    normal_style.font.color.rgb = COLOR_DARK

    # --- HEADER TABLE ---
    header_table = doc.add_table(rows=1, cols=2)
    header_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    header_table.autofit = False
    header_table.columns[0].width = Inches(4.8)
    header_table.columns[1].width = Inches(2.5)

    # Header Left: Brand Info
    cell_left = header_table.cell(0, 0)
    p_brand = cell_left.paragraphs[0]
    p_brand.paragraph_format.space_after = Pt(2)
    run_title = p_brand.add_run("DHRUTHI WELLNESS")
    run_title.font.size = Pt(20)
    run_title.font.bold = True
    run_title.font.color.rgb = COLOR_PRIMARY

    p_tagline = cell_left.add_paragraph()
    p_tagline.paragraph_format.space_after = Pt(4)
    run_tagline = p_tagline.add_run("Nourish. Balance. Thrive.")
    run_tagline.font.size = Pt(10)
    run_tagline.font.bold = True
    run_tagline.font.color.rgb = COLOR_PRIMARY

    p_email = cell_left.add_paragraph()
    p_email.paragraph_format.space_after = Pt(1)
    run_email = p_email.add_run("Email: dhruthiwellness@gmail.com")
    run_email.font.size = Pt(9)
    run_email.font.color.rgb = COLOR_MUTED

    p_phone = cell_left.add_paragraph()
    p_phone.paragraph_format.space_after = Pt(1)
    run_phone = p_phone.add_run("Phone / WhatsApp: +91 86889 63230 | +91 90524 29208")
    run_phone.font.size = Pt(9)
    run_phone.font.color.rgb = COLOR_MUTED

    p_addr = cell_left.add_paragraph()
    p_addr.paragraph_format.space_after = Pt(0)
    run_addr = p_addr.add_run("Address: 4-94, Lunani Nagar, Eluru, Andhra Pradesh, 534005, India.")
    run_addr.font.size = Pt(9)
    run_addr.font.color.rgb = COLOR_MUTED

    # Header Right: Receipt Info
    cell_right = header_table.cell(0, 1)
    p_rec = cell_right.paragraphs[0]
    p_rec.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p_rec.paragraph_format.space_after = Pt(2)
    run_rec = p_rec.add_run("PAYMENT RECEIPT")
    run_rec.font.size = Pt(16)
    run_rec.font.bold = True
    run_rec.font.color.rgb = COLOR_PRIMARY

    p_rno = cell_right.add_paragraph()
    p_rno.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p_rno.paragraph_format.space_after = Pt(1)
    r_rno = p_rno.add_run(f"Receipt No: {receipt_no}")
    r_rno.font.size = Pt(9)
    r_rno.font.bold = True

    p_rdate = cell_right.add_paragraph()
    p_rdate.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p_rdate.paragraph_format.space_after = Pt(6)
    r_rdate = p_rdate.add_run(f"Date: {payment_date}")
    r_rdate.font.size = Pt(9)

    p_badge = cell_right.add_paragraph()
    p_badge.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run_badge = p_badge.add_run("  ✓ PAYMENT SUCCESSFUL  ")
    run_badge.font.size = Pt(9)
    run_badge.font.bold = True
    run_badge.font.color.rgb = COLOR_SUCCESS

    doc.add_paragraph().paragraph_format.space_after = Pt(6)

    # --- CLIENT & PAYMENT DETAILS TABLE ---
    info_table = doc.add_table(rows=1, cols=2)
    info_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    info_table.autofit = False
    info_table.columns[0].width = Inches(3.6)
    info_table.columns[1].width = Inches(3.7)

    c0 = info_table.cell(0, 0)
    c1 = info_table.cell(0, 1)
    set_cell_background(c0, "F8F9F6")
    set_cell_background(c1, "F8F9F6")
    set_cell_margins(c0, top=140, bottom=140, left=180, right=180)
    set_cell_margins(c1, top=140, bottom=140, left=180, right=180)

    # Left: Received From
    p_c0 = c0.paragraphs[0]
    p_c0.paragraph_format.space_after = Pt(2)
    r_h0 = p_c0.add_run("RECEIVED FROM")
    r_h0.font.size = Pt(9)
    r_h0.font.bold = True
    r_h0.font.color.rgb = COLOR_PRIMARY

    p_name = c0.add_paragraph()
    p_name.paragraph_format.space_after = Pt(1)
    p_name.add_run("Name: ").bold = True
    p_name.add_run(client_name)

    p_phone = c0.add_paragraph()
    p_phone.paragraph_format.space_after = Pt(1)
    p_phone.add_run("Contact: ").bold = True
    p_phone.add_run(client_phone)

    p_email = c0.add_paragraph()
    p_email.paragraph_format.space_after = Pt(0)
    p_email.add_run("Email: ").bold = True
    p_email.add_run(client_email)

    # Right: Payment Details
    p_c1 = c1.paragraphs[0]
    p_c1.paragraph_format.space_after = Pt(2)
    r_h1 = p_c1.add_run("PAYMENT DETAILS")
    r_h1.font.size = Pt(9)
    r_h1.font.bold = True
    r_h1.font.color.rgb = COLOR_PRIMARY

    p_mode = c1.add_paragraph()
    p_mode.paragraph_format.space_after = Pt(1)
    p_mode.add_run("Mode of Payment: ").bold = True
    p_mode.add_run(payment_method)

    p_txn = c1.add_paragraph()
    p_txn.paragraph_format.space_after = Pt(1)
    p_txn.add_run("Transaction ID: ").bold = True
    p_txn.add_run(transaction_id)

    p_stat = c1.add_paragraph()
    p_stat.paragraph_format.space_after = Pt(0)
    p_stat.add_run("Status: ").bold = True
    p_stat.add_run("Completed & Verified")

    doc.add_paragraph().paragraph_format.space_after = Pt(10)

    # --- SERVICE SUMMARY TABLE ---
    items_table = doc.add_table(rows=3, cols=3)
    items_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    items_table.autofit = False
    items_table.columns[0].width = Inches(0.6)
    items_table.columns[1].width = Inches(4.9)
    items_table.columns[2].width = Inches(1.8)

    # Header Row
    hdr_cells = items_table.rows[0].cells
    headers = ["S.No", "Description / Package Name", "Amount"]
    for idx, name in enumerate(headers):
        set_cell_background(hdr_cells[idx], "3B5336")
        set_cell_margins(hdr_cells[idx], top=120, bottom=120, left=140, right=140)
        p = hdr_cells[idx].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT if idx == 2 else WD_ALIGN_PARAGRAPH.LEFT
        r = p.add_run(name)
        r.font.bold = True
        r.font.size = Pt(10)
        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    # Data Row
    row1_cells = items_table.rows[1].cells
    set_cell_margins(row1_cells[0], top=120, bottom=120, left=140, right=140)
    set_cell_margins(row1_cells[1], top=120, bottom=120, left=140, right=140)
    set_cell_margins(row1_cells[2], top=120, bottom=120, left=140, right=140)

    p_sno = row1_cells[0].paragraphs[0]
    p_sno.add_run("1")

    p_plan = row1_cells[1].paragraphs[0]
    p_plan.paragraph_format.space_after = Pt(2)
    r_plantitle = p_plan.add_run(plan_name)
    r_plantitle.font.bold = True

    p_sub = row1_cells[1].add_paragraph()
    p_sub.paragraph_format.space_after = Pt(0)
    r_sub = p_sub.add_run("Includes personalized diet consultation, progress tracking, and weekly check-ins.")
    r_sub.font.size = Pt(8.5)
    r_sub.font.color.rgb = COLOR_MUTED

    p_amt = row1_cells[2].paragraphs[0]
    p_amt.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r_amt = p_amt.add_run(f"INR {amount_paid:,.2f}")
    r_amt.font.bold = True
    r_amt.font.color.rgb = COLOR_PRIMARY

    # Total Row
    row2_cells = items_table.rows[2].cells
    set_cell_background(row2_cells[1], "F8F9F6")
    set_cell_background(row2_cells[2], "F8F9F6")
    set_cell_margins(row2_cells[1], top=120, bottom=120, left=140, right=140)
    set_cell_margins(row2_cells[2], top=120, bottom=120, left=140, right=140)

    p_tot_lbl = row2_cells[1].paragraphs[0]
    p_tot_lbl.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r_tot_lbl = p_tot_lbl.add_run("Total Amount Paid")
    r_tot_lbl.font.bold = True
    r_tot_lbl.font.size = Pt(10.5)

    p_tot_val = row2_cells[2].paragraphs[0]
    p_tot_val.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r_tot_val = p_tot_val.add_run(f"INR {amount_paid:,.2f}")
    r_tot_val.font.bold = True
    r_tot_val.font.size = Pt(11)
    r_tot_val.font.color.rgb = COLOR_PRIMARY

    doc.add_paragraph().paragraph_format.space_after = Pt(12)

    # --- NOTES & RIGHT SIGNATURE SECTION TABLE ---
    bottom_table = doc.add_table(rows=1, cols=2)
    bottom_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    bottom_table.autofit = False
    bottom_table.columns[0].width = Inches(4.8)
    bottom_table.columns[1].width = Inches(2.5)

    cell_note = bottom_table.cell(0, 0)
    cell_auth = bottom_table.cell(0, 1)

    p_note_lbl = cell_note.paragraphs[0]
    p_note_lbl.paragraph_format.space_after = Pt(2)
    r_nl = p_note_lbl.add_run("Note:")
    r_nl.font.bold = True

    p_note_val = cell_note.add_paragraph()
    r_nv = p_note_val.add_run(notes)
    r_nv.font.size = Pt(9)
    r_nv.font.color.rgb = COLOR_MUTED

    # Right: Signature Block
    p_auth_title = cell_auth.paragraphs[0]
    p_auth_title.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p_auth_title.paragraph_format.space_after = Pt(6)
    r_at = p_auth_title.add_run("For Dhruthi Wellness")
    r_at.font.bold = True
    r_at.font.size = Pt(10)
    r_at.font.color.rgb = COLOR_PRIMARY

    # Add Signature Image if available
    sig_path_png = os.path.join(os.path.dirname(os.path.dirname(os.path.abspath(__file__))), "Images", "akhi_sign_transparent.png")
    if os.path.exists(sig_path_png):
        p_sig = cell_auth.add_paragraph()
        p_sig.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p_sig.paragraph_format.space_after = Pt(12)
        run_sig = p_sig.add_run()
        run_sig.add_picture(sig_path_png, width=Inches(1.5))

    p_sig_lbl = cell_auth.add_paragraph()
    p_sig_lbl.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r_sl = p_sig_lbl.add_run("Authorized Signatory")
    r_sl.font.italic = True
    r_sl.font.size = Pt(9)
    r_sl.font.color.rgb = COLOR_MUTED

    doc.add_paragraph().paragraph_format.space_after = Pt(20)

    # --- FOOTER DECLARATION ---
    p_ftr = doc.add_paragraph()
    p_ftr.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_ftr = p_ftr.add_run("This is an official computer-generated receipt issued by Dhruthi Wellness.")
    r_ftr.font.size = Pt(8.5)
    r_ftr.font.color.rgb = COLOR_MUTED

    # Save Word document in docs directory
    docs_dir = os.path.join(os.path.dirname(os.path.dirname(os.path.abspath(__file__))), "docs")
    os.makedirs(docs_dir, exist_ok=True)
    output_path = os.path.join(docs_dir, output_filename)
    doc.save(output_path)
    print(f"Successfully generated editable Word receipt: {output_path}")
    return output_path

if __name__ == "__main__":
    create_receipt_docx()
