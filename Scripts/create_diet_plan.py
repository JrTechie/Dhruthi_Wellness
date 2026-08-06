import os
import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls

def set_cell_background(cell, fill_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=100, bottom=100, left=120, right=120):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = parse_xml(
        f'<w:tcMar {nsdecls("w")}>'
        f'<w:top w:w="{top}" w:type="dxa"/>'
        f'<w:bottom w:w="{bottom}" w:type="dxa"/>'
        f'<w:left w:w="{left}" w:type="dxa"/>'
        f'<w:right w:w="{right}" w:type="dxa"/>'
        f'</w:tcMar>'
    )
    tcPr.append(tcMar)

def set_table_borders(table, color="CBD5E1", sz="4"):
    tblPr = table._tbl.tblPr
    tblBorders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        f'<w:top w:val="single" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        f'<w:bottom w:val="single" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        f'<w:left w:val="single" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        f'<w:right w:val="single" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        f'<w:insideH w:val="single" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        f'<w:insideV w:val="single" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        f'</w:tblBorders>'
    )
    tblPr.append(tblBorders)

def make_row_cant_split(row):
    trPr = row._tr.get_or_add_trPr()
    trPr.append(parse_xml(f'<w:cantSplit {nsdecls("w")}/>'))

def make_row_header(row):
    trPr = row._tr.get_or_add_trPr()
    trPr.append(parse_xml(f'<w:tblHeader {nsdecls("w")}/>'))

def create_document():
    doc = docx.Document()
    
    # Configure landscape page layout (11.0 in x 8.5 in)
    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = Inches(11.0)
    section.page_height = Inches(8.5)
    section.top_margin = Inches(0.4)
    section.bottom_margin = Inches(0.4)
    section.left_margin = Inches(0.5)
    section.right_margin = Inches(0.5)
    
    # Colors Palette
    COLOR_PRIMARY = RGBColor(15, 76, 92)     # #0F4C5C (Teal)
    COLOR_SECONDARY = RGBColor(157, 23, 77)  # #9D174D (Magenta/Berry)
    COLOR_TEXT = RGBColor(34, 34, 34)        # #222222
    COLOR_MUTED = RGBColor(100, 116, 139)    # #64748B
    COLOR_GREEN = RGBColor(16, 185, 129)     # Emerald Green
    
    HEX_HEADER_BG = "0F4C5C"
    HEX_SUBHEADER_BG = "F1F5F9"
    HEX_ACCENT_BG = "FDF2F8"
    HEX_CARD_BG = "F8FAFC"
    
    logo_path = r"l:\Developer\nutriflow\Images\Logo_Dhruthi_Wellness.png"

    # --- TOP HEADER BAR ---
    header_table = doc.add_table(rows=1, cols=2)
    header_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    header_table.autofit = False
    header_table.columns[0].width = Inches(7.5)
    header_table.columns[1].width = Inches(2.5)
    
    cell_left = header_table.rows[0].cells[0]
    cell_right = header_table.rows[0].cells[1]
    
    p_title = cell_left.paragraphs[0]
    p_title.paragraph_format.space_before = Pt(0)
    p_title.paragraph_format.space_after = Pt(2)
    r_title = p_title.add_run("{user name’s} Diet Plan: Week 1- Week 4")
    r_title.font.name = 'Arial'
    r_title.font.size = Pt(22)
    r_title.font.bold = True
    r_title.font.color.rgb = COLOR_PRIMARY
    
    p_sub = cell_left.add_paragraph()
    p_sub.paragraph_format.space_after = Pt(0)
    r_sub = p_sub.add_run("DHRUTHI WELLNESS • CUSTOMIZED NUTRITION & FERTILITY CARE")
    r_sub.font.name = 'Arial'
    r_sub.font.size = Pt(9)
    r_sub.font.bold = True
    r_sub.font.color.rgb = COLOR_SECONDARY
    
    p_logo = cell_right.paragraphs[0]
    p_logo.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p_logo.paragraph_format.space_after = Pt(0)
    if os.path.exists(logo_path):
        p_logo.add_run().add_picture(logo_path, width=Inches(2.2))
    else:
        r_logo = p_logo.add_run("DHRUTHI WELLNESS")
        r_logo.font.name = 'Arial'
        r_logo.font.size = Pt(16)
        r_logo.font.bold = True
        r_logo.font.color.rgb = COLOR_PRIMARY
        
    doc.add_paragraph().paragraph_format.space_after = Pt(4)
    
    # --- METRICS & SUMMARY CARD ---
    metrics_table = doc.add_table(rows=4, cols=1)
    metrics_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    metrics_table.autofit = False
    metrics_table.columns[0].width = Inches(10.0)
    set_table_borders(metrics_table, color="CBD5E1", sz="6")
    
    metrics_data = [
        [("History: ", True, COLOR_TEXT), ("32 years | Trying to Conceive | PCOS, Thyroid", False, COLOR_TEXT)],
        [("BMI: ", True, COLOR_TEXT), ("65 Kgs | 5 feet 4 inches | BMI is 23.5 – ", False, COLOR_TEXT), ("Normal", True, COLOR_GREEN), ("/ Low/ High", False, COLOR_MUTED)],
        [("Recommended Calories & Macro Split: ", True, COLOR_TEXT), ("Protein: 65–70g | Carbohydrates: 130g | Healthy Fats: 35g", False, COLOR_PRIMARY)],
        [("Weight loss Goal: ", True, COLOR_SECONDARY), ("0.5–0.75 kg per week for first 3 months → ", False, COLOR_TEXT), ("5–10% weight loss can improve your fertility & cycles 📅", True, COLOR_SECONDARY)]
    ]
    
    for i, row_data in enumerate(metrics_data):
        cell = metrics_table.rows[i].cells[0]
        set_cell_background(cell, HEX_CARD_BG)
        set_cell_margins(cell, top=50, bottom=50, left=120, right=120)
        p = cell.paragraphs[0]
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after = Pt(1)
        p.paragraph_format.line_spacing = Pt(13)
        
        for text, bold, color in row_data:
            r = p.add_run(text)
            r.font.name = 'Arial'
            r.font.size = Pt(10)
            r.font.bold = bold
            r.font.color.rgb = color
            
    doc.add_paragraph().paragraph_format.space_after = Pt(6)
    
    # --- WEEK 1-4 INTERVENTIONS SECTION ---
    p_sec1 = doc.add_paragraph()
    p_sec1.paragraph_format.space_before = Pt(2)
    p_sec1.paragraph_format.space_after = Pt(3)
    r_sec1 = p_sec1.add_run("Week 1-4 Interventions")
    r_sec1.font.name = 'Arial'
    r_sec1.font.size = Pt(13)
    r_sec1.font.bold = True
    r_sec1.font.color.rgb = COLOR_PRIMARY
    
    interventions_table = doc.add_table(rows=3, cols=2)
    interventions_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    interventions_table.autofit = False
    interventions_table.columns[0].width = Inches(3.2)
    interventions_table.columns[1].width = Inches(6.8)
    set_table_borders(interventions_table, color="CBD5E1", sz="6")
    
    # Row 0: Goals Header
    r0_c0 = interventions_table.rows[0].cells[0]
    r0_c1 = interventions_table.rows[0].cells[1]
    set_cell_background(r0_c0, HEX_ACCENT_BG)
    set_cell_background(r0_c1, HEX_ACCENT_BG)
    set_cell_margins(r0_c0, top=60, bottom=60, left=120, right=120)
    set_cell_margins(r0_c1, top=60, bottom=60, left=120, right=120)
    
    p = r0_c0.paragraphs[0]
    r = p.add_run("Goals")
    r.font.name = 'Arial'; r.font.size = Pt(10.5); r.font.bold = True; r.font.color.rgb = COLOR_SECONDARY
    
    p = r0_c1.paragraphs[0]
    r = p.add_run("Intervention Details & Recommendations")
    r.font.name = 'Arial'; r.font.size = Pt(10.5); r.font.bold = True; r.font.color.rgb = COLOR_SECONDARY
    
    # Row 1: Blood Test
    r1_c0 = interventions_table.rows[1].cells[0]
    r1_c1 = interventions_table.rows[1].cells[1]
    set_cell_margins(r1_c0, top=60, bottom=60, left=120, right=120)
    set_cell_margins(r1_c1, top=60, bottom=60, left=120, right=120)
    
    p = r1_c0.paragraphs[0]
    r = p.add_run("Get a Comprehensive Blood Test")
    r.font.name = 'Arial'; r.font.size = Pt(10); r.font.bold = True; r.font.color.rgb = COLOR_SECONDARY
    
    p = r1_c1.paragraphs[0]
    r = p.add_run("Understand deficiencies & hormonal imbalances. ")
    r.font.name = 'Arial'; r.font.size = Pt(10); r.font.color.rgb = COLOR_TEXT
    r_link = p.add_run("LINK to BOOK.")
    r_link.font.name = 'Arial'; r_link.font.size = Pt(10); r_link.font.bold = True; r_link.font.underline = True; r_link.font.color.rgb = COLOR_PRIMARY
    
    # Row 2: Steps & Yoga
    r2_c0 = interventions_table.rows[2].cells[0]
    r2_c1 = interventions_table.rows[2].cells[1]
    set_cell_margins(r2_c0, top=60, bottom=60, left=120, right=120)
    set_cell_margins(r2_c1, top=60, bottom=60, left=120, right=120)
    
    p = r2_c0.paragraphs[0]
    r = p.add_run("Daily 7,000 Steps & Yoga 5x per week")
    r.font.name = 'Arial'; r.font.size = Pt(10); r.font.bold = True; r.font.color.rgb = COLOR_SECONDARY
    
    p = r2_c1.paragraphs[0]
    r = p.add_run("Recommend to practice ")
    r.font.name = 'Arial'; r.font.size = Pt(10); r.font.color.rgb = COLOR_TEXT
    r_bold = p.add_run("Live Yoga for Fertility")
    r_bold.font.name = 'Arial'; r_bold.font.size = Pt(10); r_bold.font.bold = True; r_bold.font.color.rgb = COLOR_TEXT
    r_rest = p.add_run(" 3–5 times a week. Walk for 5–10 mins after each meal.")
    r_rest.font.name = 'Arial'; r_rest.font.size = Pt(10); r_rest.font.color.rgb = COLOR_TEXT

    doc.add_paragraph().paragraph_format.space_after = Pt(8)

    # --- YOUR MEAL PLAN MATRIX TABLE ---
    p_sec2 = doc.add_paragraph()
    p_sec2.paragraph_format.space_before = Pt(2)
    p_sec2.paragraph_format.space_after = Pt(3)
    r_sec2 = p_sec2.add_run("YOUR MEAL PLAN 🍽️")
    r_sec2.font.name = 'Arial'
    r_sec2.font.size = Pt(13)
    r_sec2.font.bold = True
    r_sec2.font.color.rgb = COLOR_PRIMARY

    col_widths = [Inches(1.6), Inches(1.2), Inches(1.2), Inches(1.2), Inches(1.2), Inches(1.2), Inches(1.2), Inches(1.2)]
    
    meal_table = doc.add_table(rows=7, cols=8)
    meal_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    meal_table.autofit = False
    set_table_borders(meal_table, color="CBD5E1", sz="4")
    
    for i, w in enumerate(col_widths):
        meal_table.columns[i].width = w

    headers = ["", "Monday", "Tuesday", "Wednesday", "Thursday", "Friday", "Saturday", "Sunday"]
    
    # Header Row
    hdr_row = meal_table.rows[0]
    make_row_cant_split(hdr_row)
    make_row_header(hdr_row)
    for col_idx, text in enumerate(headers):
        cell = hdr_row.cells[col_idx]
        set_cell_background(cell, HEX_HEADER_BG)
        set_cell_margins(cell, top=80, bottom=80, left=40, right=40)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(text)
        r.font.name = 'Arial'; r.font.size = Pt(10.5); r.font.bold = True; r.font.color.rgb = RGBColor(255, 255, 255)

    # Helper for building meal rows
    def setup_left_meal_cell(row, category_title, time_sub, macro_sub):
        cell_0 = row.cells[0]
        set_cell_background(cell_0, HEX_SUBHEADER_BG)
        set_cell_margins(cell_0, top=60, bottom=60, left=50, right=50)
        p0 = cell_0.paragraphs[0]
        p0.paragraph_format.space_after = Pt(2)
        r0 = p0.add_run(category_title)
        r0.font.name = 'Arial'; r0.font.size = Pt(10); r0.font.bold = True; r0.font.color.rgb = COLOR_SECONDARY
        
        if time_sub:
            p_time = cell_0.add_paragraph()
            p_time.paragraph_format.space_after = Pt(2)
            r_t = p_time.add_run(time_sub)
            r_t.font.name = 'Arial'; r_t.font.size = Pt(8); r_t.font.color.rgb = COLOR_MUTED
            
        if macro_sub:
            p_macro = cell_0.add_paragraph()
            p_macro.paragraph_format.space_after = Pt(0)
            r_m = p_macro.add_run(macro_sub)
            r_m.font.name = 'Arial'; r_m.font.size = Pt(7.5); r_m.font.italic = True; r_m.font.color.rgb = COLOR_PRIMARY

    # 1. Pre-Breakfast (Merged across all 7 days)
    row_1 = meal_table.rows[1]
    make_row_cant_split(row_1)
    setup_left_meal_cell(row_1, "Pre-Breakfast", "7.30 – 8.15 am", "")
    
    # Merge cells 1 through 7
    merged_pre = row_1.cells[1].merge(row_1.cells[7])
    set_cell_margins(merged_pre, top=60, bottom=60, left=80, right=80)
    p_pre = merged_pre.paragraphs[0]
    p_pre.paragraph_format.space_after = Pt(0)
    r_tag = p_pre.add_run("Fertility Detox: ")
    r_tag.font.name = 'Arial'; r_tag.font.size = Pt(9); r_tag.font.bold = True; r_tag.font.color.rgb = COLOR_SECONDARY
    r_txt = p_pre.add_run("Jeera water - 1 cup (Boil ½ tsp of cumin seeds in a cup of water for 2-3 min, strain and drink) + 5-6 soaked, peeled almond + 2 soaked walnuts")
    r_txt.font.name = 'Arial'; r_txt.font.size = Pt(9); r_txt.font.color.rgb = COLOR_TEXT

    # 2. Breakfast
    row_2 = meal_table.rows[2]
    make_row_cant_split(row_2)
    bf_macros = "45% Carbs, 20% Protein, 15% Fiber, 20% Fat\n300-350 calories"
    setup_left_meal_cell(row_2, "Breakfast", "8.30 – 9.30 am", bf_macros)
    bf_days = [
        "2 moong dal dosa + 1 tbsp coconut chutney + sambar (1 cup)",
        "Poha with added peanuts, green peas (1 cup) + ½ cup steamed sprouts",
        "Besan chilla (2 small) + green chutney + ½ cup curd",
        "1 cup ragi kanji (ragi porridge) + ½ cup sprouted moong",
        "Vegetable paneer paratha (1 medium) + green chutney + curd (½ cup)",
        "2 rava uthappam + tomato chutney + ½ cup sauteed paneer",
        "1 bowl overnight oatmeal (made with rolled oats) topped with 1 tbsp unsweetened peanut butter"
    ]
    for d_idx, text in enumerate(bf_days):
        cell = row_2.cells[d_idx + 1]
        set_cell_margins(cell, top=60, bottom=60, left=50, right=50)
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0); p.paragraph_format.line_spacing = Pt(11.5)
        r = p.add_run(text)
        r.font.name = 'Arial'; r.font.size = Pt(8.5); r.font.color.rgb = COLOR_TEXT

    # 3. Mid-Morning (Merged across all 7 days)
    row_3 = meal_table.rows[3]
    make_row_cant_split(row_3)
    setup_left_meal_cell(row_3, "Mid-Morning", "12.30 pm", "")
    merged_mm = row_3.cells[1].merge(row_3.cells[7])
    set_cell_margins(merged_mm, top=60, bottom=60, left=80, right=80)
    p_mm = merged_mm.paragraphs[0]
    p_mm.paragraph_format.space_after = Pt(0)
    r_mm = p_mm.add_run("Fruit (any seasonal) - 1 no. + 1-2 tbsp unsweetened almond butter")
    r_mm.font.name = 'Arial'; r_mm.font.size = Pt(9); r_mm.font.color.rgb = COLOR_TEXT

    # 4. Lunch
    row_4 = meal_table.rows[4]
    make_row_cant_split(row_4)
    lunch_macros = "45% Carbs, 20% Protein, 15% Fiber, 20% Fat\n400-450 calories"
    setup_left_meal_cell(row_4, "Lunch", "2:00 pm", lunch_macros)
    salad_note = "1 Cup Salad (15 mins before lunch)"
    lunch_days = [
        "Roti (1) + Dal (1 cup) + Sabzi (1 cup) + Paneer bhurji (½ cup)",
        "Roti (1) + Dal (1 cup) + Green leafy vegetable Sabzi/ Poriyal (1 cup) + Curd (½ cup)",
        "1 roti + 1 cup sundal (chana stir-fry) + mixed veg poriyal",
        "1 roti + 1 cup moringa leaf dal + cucumber raita",
        "Scrambled tofu with preferred choice of vegetables + 2 small plain chapati",
        "1 small bajra roti + 1 cup paneer curry + ½ cup curd",
        "Brown rice (½ cup) + 1 cup kootu + 1 cup sabzi/ poriyal"
    ]
    for d_idx, text in enumerate(lunch_days):
        cell = row_4.cells[d_idx + 1]
        set_cell_margins(cell, top=60, bottom=60, left=50, right=50)
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(2)
        r_s = p.add_run(salad_note + "\n")
        r_s.font.name = 'Arial'; r_s.font.size = Pt(8); r_s.font.italic = True; r_s.font.bold = True; r_s.font.color.rgb = COLOR_SECONDARY
        r_t = p.add_run(text)
        r_t.font.name = 'Arial'; r_t.font.size = Pt(8.5); r_t.font.color.rgb = COLOR_TEXT

    # 5. Snack
    row_5 = meal_table.rows[5]
    make_row_cant_split(row_5)
    setup_left_meal_cell(row_5, "Snack", "4.30 – 5.30 pm", "")
    snack_days = [
        "Makhana (¾ cup) + Green tea + Dates (2)",
        "Handful of roasted chana + Green tea",
        "Makhana (¾ cup) + Green tea + Dates (2)",
        "Steamed sprouts chaat (¾ cup) + Green tea",
        "Makhana (¾ cup) + Green tea + Dates (2)",
        "Ragi khakhra (2) + Green tea",
        "Makhana (¾ cup) + Green tea + Dates (2)"
    ]
    for d_idx, text in enumerate(snack_days):
        cell = row_5.cells[d_idx + 1]
        set_cell_margins(cell, top=60, bottom=60, left=50, right=50)
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0); p.paragraph_format.line_spacing = Pt(11.5)
        r = p.add_run(text)
        r.font.name = 'Arial'; r.font.size = Pt(8.5); r.font.color.rgb = COLOR_TEXT

    # 6. Dinner
    row_6 = meal_table.rows[6]
    make_row_cant_split(row_6)
    dinner_macros = "45% Carbs, 20% Protein, 15% Fiber, 20% Fat\n300-350 calories"
    setup_left_meal_cell(row_6, "Dinner", "7.30 – 8.30 pm", dinner_macros)
    dinner_days = [
        "Jowar roti (1) + Paneer Sabzi (1 cup)",
        "1 bowl tomato rasam + 1 small bowl thinai (foxtail millet) rice",
        "1 bowl lemon rice + curd-based kadhi (½ cup) + sautéed veggies",
        "1 millet dosa + vegetable stew (½ cup)",
        "1 cup brown rice + mixed dal + sautéed veggies",
        "1 bowl lemon rice + vegetable curry (½ cup) + sautéed veggies",
        "Grilled paneer salad (70-80 grams) + 1 multigrain roti + sabzi"
    ]
    for d_idx, text in enumerate(dinner_days):
        cell = row_6.cells[d_idx + 1]
        set_cell_margins(cell, top=60, bottom=60, left=50, right=50)
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0); p.paragraph_format.line_spacing = Pt(11.5)
        r = p.add_run(text)
        r.font.name = 'Arial'; r.font.size = Pt(8.5); r.font.color.rgb = COLOR_TEXT

    # Page Break for Guidelines
    doc.add_page_break()

    # --- PAGE 2: BEGINNER FERTILITY GUIDELINES ---
    header_table2 = doc.add_table(rows=1, cols=2)
    header_table2.alignment = WD_TABLE_ALIGNMENT.CENTER
    header_table2.autofit = False
    header_table2.columns[0].width = Inches(7.5)
    header_table2.columns[1].width = Inches(2.5)
    
    cell_l2 = header_table2.rows[0].cells[0]
    cell_r2 = header_table2.rows[0].cells[1]
    
    p2_t = cell_l2.paragraphs[0]
    p2_t.paragraph_format.space_before = Pt(0)
    p2_t.paragraph_format.space_after = Pt(2)
    r2_t = p2_t.add_run("{user name’s} Diet Plan: Week 1- Week 4")
    r2_t.font.name = 'Arial'; r2_t.font.size = Pt(18); r2_t.font.bold = True; r2_t.font.color.rgb = COLOR_PRIMARY
    
    p2_logo = cell_r2.paragraphs[0]
    p2_logo.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    if os.path.exists(logo_path):
        p2_logo.add_run().add_picture(logo_path, width=Inches(1.8))
        
    doc.add_paragraph().paragraph_format.space_after = Pt(4)

    # Beginner Fertility Guidelines Header
    p_fg = doc.add_paragraph()
    p_fg.paragraph_format.space_before = Pt(4)
    p_fg.paragraph_format.space_after = Pt(2)
    r_fg = p_fg.add_run("🧘 Beginner Fertility Guidelines")
    r_fg.font.name = 'Arial'; r_fg.font.size = Pt(14); r_fg.font.bold = True; r_fg.font.color.rgb = COLOR_PRIMARY

    p_fg_sub = doc.add_paragraph()
    p_fg_sub.paragraph_format.space_after = Pt(8)
    r_fg_sub = p_fg_sub.add_run("To support hormonal balance, egg quality, and ovulation, it’s essential to limit or avoid the following foods:")
    r_fg_sub.font.name = 'Arial'; r_fg_sub.font.size = Pt(10.5); r_fg_sub.font.color.rgb = COLOR_TEXT

    guidelines = [
        ("1️⃣ For Improved Egg Quality → Cut out Processed & Packaged Foods",
         "Why?", " High in trans fats, preservatives, and artificial additives, which can cause hormonal imbalances and inflammation, affecting egg quality.",
         "Examples:", " Instant noodles, chips, biscuits, store-bought bakery items, frozen ready-to-eat meals."),
        
        ("2️⃣ For Estrogen Balance & Ovulation → Avoid Excess Caffeine & Sugary Beverages ☕",
         "Why?", " High caffeine intake (>200 mg/day) can interfere with estrogen balance and ovulation. Sugary drinks cause insulin resistance, worsening symptoms.",
         "Recommendation:", " Limit coffee to 1-2 cups per day. Avoid colas, energy drinks, flavored teas, and packaged fruit juices."),
        
        ("3️⃣ For Improved Ovulation → Avoid Refined Carbohydrates & Sugars 🧂",
         "Why?", " White rice, maida, and refined sugars lead to spikes in insulin, which can worsen symptoms and delay ovulation.",
         "Examples:", " Avoid white bread, pastries, candies, and refined cereals. Choose whole grains like millets, brown rice, and whole wheat."),
        
        ("4️⃣ For Hormonal Balance → Avoid Fried & High-Sodium Foods",
         "Why?", " Excess salt and deep-fried foods cause water retention, bloating, and disrupt hormonal balance. Trans fats are also linked to reduced fertility.",
         "Examples:", " Avoid fast food, fried snacks (pakoras, samosas, chips), and restaurant-made deep-fried foods.")
    ]

    for title, label1, text1, label2, text2 in guidelines:
        g_table = doc.add_table(rows=1, cols=1)
        g_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        g_table.autofit = False
        g_table.columns[0].width = Inches(10.0)
        set_table_borders(g_table, color="E2E8F0", sz="4")
        
        c = g_table.rows[0].cells[0]
        set_cell_background(c, HEX_CARD_BG)
        set_cell_margins(c, top=60, bottom=60, left=120, right=120)
        
        p_t = c.paragraphs[0]
        p_t.paragraph_format.space_after = Pt(2)
        r = p_t.add_run(title)
        r.font.name = 'Arial'; r.font.size = Pt(11); r.font.bold = True; r.font.color.rgb = COLOR_SECONDARY
        
        p_b1 = c.add_paragraph()
        p_b1.paragraph_format.space_after = Pt(2)
        r_lbl1 = p_b1.add_run(label1)
        r_lbl1.font.name = 'Arial'; r_lbl1.font.size = Pt(10); r_lbl1.font.bold = True; r_lbl1.font.color.rgb = COLOR_TEXT
        r_txt1 = p_b1.add_run(text1)
        r_txt1.font.name = 'Arial'; r_txt1.font.size = Pt(10); r_txt1.font.color.rgb = COLOR_TEXT
        
        p_b2 = c.add_paragraph()
        p_b2.paragraph_format.space_after = Pt(1)
        r_lbl2 = p_b2.add_run(label2)
        r_lbl2.font.name = 'Arial'; r_lbl2.font.size = Pt(10); r_lbl2.font.bold = True; r_lbl2.font.italic = True; r_lbl2.font.color.rgb = COLOR_TEXT
        r_txt2 = p_b2.add_run(text2)
        r_txt2.font.name = 'Arial'; r_txt2.font.size = Pt(10); r_txt2.font.italic = True; r_txt2.font.color.rgb = COLOR_TEXT
        
        doc.add_paragraph().paragraph_format.space_after = Pt(3)

    # What to do instead
    p_wtd = doc.add_paragraph()
    p_wtd.paragraph_format.space_before = Pt(4)
    p_wtd.paragraph_format.space_after = Pt(3)
    r_wtd = p_wtd.add_run("📌 What to Do Instead?")
    r_wtd.font.name = 'Arial'; r_wtd.font.size = Pt(12); r_wtd.font.bold = True; r_wtd.font.color.rgb = COLOR_PRIMARY

    instead_items = [
        "Opt for whole, home-cooked, nutrient-dense meals",
        "Drink herbal teas & infused water instead of sugary drinks",
        "Choose unprocessed, fiber-rich carbs (quinoa, millets, steel-cut oats)",
        "Swap deep-fried snacks for roasted makhana, nuts, or homemade energy bars"
    ]

    for item in instead_items:
        p_item = doc.add_paragraph()
        p_item.paragraph_format.left_indent = Inches(0.2)
        p_item.paragraph_format.space_after = Pt(2)
        r_chk = p_item.add_run("✔  ")
        r_chk.font.name = 'Arial'; r_chk.font.size = Pt(10); r_chk.font.bold = True; r_chk.font.color.rgb = COLOR_GREEN
        r_txt = p_item.add_run(item)
        r_txt.font.name = 'Arial'; r_txt.font.size = Pt(10); r_txt.font.color.rgb = COLOR_TEXT

    doc.add_paragraph().paragraph_format.space_after = Pt(6)

    # Next Steps
    p_ns = doc.add_paragraph()
    p_ns.paragraph_format.space_before = Pt(4)
    p_ns.paragraph_format.space_after = Pt(3)
    r_ns = p_ns.add_run("Next Steps")
    r_ns.font.name = 'Arial'; r_ns.font.size = Pt(12); r_ns.font.bold = True; r_ns.font.color.rgb = COLOR_PRIMARY

    steps = [
        ("📸 Daily Meal Tracking: ", "Share pictures of your meals daily on WhatsApp for feedback and guidance."),
        ("📞 Check-in Call (Day 7): ", "In 7 days, we’ll have a one-on-one check-in call to review your progress, discuss your blood test reports, make adjustments to your plan."),
        ("🌸 Remember: ", "Your consistency and engagement are key to improving your hormonal balance and overall fertility health. Let’s take this journey together—one healthy meal at a time! 🌱")
    ]

    for label, desc in steps:
        p_st = doc.add_paragraph()
        p_st.paragraph_format.left_indent = Inches(0.2)
        p_st.paragraph_format.space_after = Pt(3)
        r_lbl = p_st.add_run(label)
        r_lbl.font.name = 'Arial'; r_lbl.font.size = Pt(10); r_lbl.font.bold = True; r_lbl.font.color.rgb = COLOR_TEXT
        r_desc = p_st.add_run(desc)
        r_desc.font.name = 'Arial'; r_desc.font.size = Pt(10); r_desc.font.color.rgb = COLOR_TEXT

    p_footer = doc.add_paragraph()
    p_footer.paragraph_format.space_before = Pt(6)
    p_footer.paragraph_format.space_after = Pt(0)
    r_foot = p_footer.add_run("If you have any questions, feel free to reach out on WhatsApp! 😊")
    r_foot.font.name = 'Arial'; r_foot.font.size = Pt(10.5); r_foot.font.bold = True; r_foot.font.italic = True; r_foot.font.color.rgb = COLOR_PRIMARY

    output_filename = "Dhruthi_Wellness_Diet_Plan_Template.docx"
    doc.save(output_filename)
    print(f"Successfully generated {output_filename}")

if __name__ == "__main__":
    create_document()
